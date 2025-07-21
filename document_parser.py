import os
import fitz  # PyMuPDF
from docx import Document
from typing import List, Dict, Any
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentParser:
    """Open-source document parser for PDF and Word documents"""
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.doc']
    
    def parse_pdf(self, file_path: str) -> str:
        """Parse PDF file and extract text content"""
        try:
            doc = fitz.open(file_path)
            text_content = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text_content += page.get_text()
                text_content += "\n\n"  # Add page separator
            
            doc.close()
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            return ""
    
    def parse_docx(self, file_path: str) -> str:
        """Parse DOCX file and extract text content"""
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return ""
    
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """Parse a document and return its content with metadata"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        file_name = os.path.basename(file_path)
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        logger.info(f"Parsing document: {file_name}")
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text_content = self.parse_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            text_content = self.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        if not text_content:
            logger.warning(f"No text content extracted from {file_name}")
        
        return {
            'file_path': file_path,
            'file_name': file_name,
            'file_extension': file_extension,
            'text_content': text_content,
            'content_length': len(text_content),
            'word_count': len(text_content.split()) if text_content else 0
        }
    
    def parse_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """Parse all supported documents in a directory"""
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        parsed_documents = []
        
        for file_name in os.listdir(directory_path):
            file_path = os.path.join(directory_path, file_name)
            
            if os.path.isfile(file_path):
                file_extension = os.path.splitext(file_name)[1].lower()
                
                if file_extension in self.supported_extensions:
                    try:
                        parsed_doc = self.parse_document(file_path)
                        parsed_documents.append(parsed_doc)
                    except Exception as e:
                        logger.error(f"Failed to parse {file_name}: {e}")
                        continue
        
        logger.info(f"Successfully parsed {len(parsed_documents)} documents from {directory_path}")
        return parsed_documents
    
    def get_supported_extensions(self) -> List[str]:
        """Return list of supported file extensions"""
        return self.supported_extensions.copy()